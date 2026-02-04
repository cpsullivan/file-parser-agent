"""Word document parser using python-docx."""

from typing import Dict, Any, List
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from .base_parser import BaseParser


class WordParser(BaseParser):
    """
    Parser for Word documents (.docx, .doc).
    
    Extracts:
    - Paragraph text with style information
    - Tables with row/column structure
    - Document metadata (title, author, dates)
    - Embedded object counts (images, OLE objects)
    """
    
    SUPPORTED_EXTENSIONS = ['.docx', '.doc']
    FILE_TYPE = 'word'
    
    def parse(self) -> Dict[str, Any]:
        """
        Parse Word document and extract content.
        
        Returns:
            Dictionary containing:
            - filename, file_type, parsed_at (base fields)
            - total_paragraphs, total_tables, image_count
            - metadata: Document metadata dict
            - paragraphs: List of paragraph objects
            - tables: List of table objects with data
        """
        output = self.get_base_output()
        
        try:
            doc = Document(self.filepath)
            
            # Extract metadata
            output['metadata'] = self._extract_metadata(doc)
            
            # Extract paragraphs
            paragraphs = self._extract_paragraphs(doc)
            output['total_paragraphs'] = len(paragraphs)
            output['paragraphs'] = paragraphs
            
            # Extract tables
            tables = self._extract_tables(doc)
            output['total_tables'] = len(tables)
            output['tables'] = tables
            
            # Count embedded objects
            image_count, embedded_count = self._count_embedded_objects(doc)
            output['image_count'] = image_count
            output['embedded_object_count'] = embedded_count
            
        except Exception as e:
            output['error'] = f"Word parsing error: {str(e)}"
            output['total_paragraphs'] = 0
            output['total_tables'] = 0
            output['image_count'] = 0
            output['embedded_object_count'] = 0
            output['paragraphs'] = []
            output['tables'] = []
            output['metadata'] = {}
        
        return output
    
    def _extract_metadata(self, doc: Document) -> Dict[str, str]:
        """
        Extract document metadata from Word file.
        
        Args:
            doc: python-docx Document instance
            
        Returns:
            Dictionary with title, author, subject, created, modified
        """
        props = doc.core_properties
        
        return {
            'title': self._safe_str(props.title),
            'author': self._safe_str(props.author),
            'subject': self._safe_str(props.subject),
            'created': self._safe_str(props.created) if props.created else '',
            'modified': self._safe_str(props.modified) if props.modified else ''
        }
    
    def _extract_paragraphs(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extract paragraphs with style information.
        
        Args:
            doc: python-docx Document instance
            
        Returns:
            List of paragraph dictionaries with text and style
        """
        paragraphs = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Only include non-empty paragraphs
                paragraphs.append({
                    'text': text,
                    'style': para.style.name if para.style else 'Normal'
                })
        
        return paragraphs
    
    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """
        Extract tables with row/column data.
        
        Args:
            doc: python-docx Document instance
            
        Returns:
            List of table dictionaries with dimensions and data
        """
        tables = []
        
        for i, table in enumerate(doc.tables, 1):
            table_data = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    # Get cell text, handling merged cells
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            
            tables.append({
                'table_number': i,
                'rows': len(table.rows),
                'columns': len(table.columns),
                'data': table_data
            })
        
        return tables
    
    def _count_embedded_objects(self, doc: Document) -> tuple:
        """
        Count embedded images and OLE objects.
        
        Args:
            doc: python-docx Document instance
            
        Returns:
            Tuple of (image_count, ole_object_count)
        """
        image_count = 0
        ole_count = 0
        
        try:
            for rel in doc.part.rels.values():
                if rel.reltype == RT.IMAGE:
                    image_count += 1
                elif 'oleObject' in str(rel.reltype).lower():
                    ole_count += 1
        except Exception:
            pass  # Some documents may not have relationship info
        
        return image_count, ole_count
